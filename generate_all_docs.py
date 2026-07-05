import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Define projects metadata and content
projects = [
    {
        "id": "cns",
        "title": "Cyber Nervous System (CNS)",
        "docs_dir": "D:/projects/rishi/cns-mvp/docs",
        "filename": "cns_technical_documentation",
        "content": """# Cyber Nervous System (CNS) — Technical Documentation
*Autonomous Code Auditing & Software Digital Twin Platform*

This document provides a comprehensive technical overview of the **Cyber Nervous System (CNS)** project located at `D:/projects/rishi/cns-mvp`. It covers the system architecture, code parser pipeline, API layers, Neo4j schema design, and operational instructions.

---

## 1. Executive Summary
The **Cyber Nervous System (CNS)** is a software security "Digital Twin" application. It recursively ingests target source codebases, parses them down to Abstract Syntax Tree (AST) definitions, constructs a relational dependency graph in Neo4j, matches libraries against real-world vulnerability databases (OSV.dev), and applies a Large Language Model (Google Gemini 1.5 Flash) to analyze the resulting graph database for potential security design flaws.

---

## 2. System Architecture

The project is structured into three primary packages:
* **`frontend/`**: Vite + React 19 single-page application representing the dashboard UI.
* **`backend/`**: FastAPI implementation routing requests, executing background ingestion, and querying Neo4j.
* **`ingestor/`**: Core parser engines responsible for AST compilation (`parser.py`) and vulnerability API lookup (`vulnerability_scanner.py`).

### A. Code Ingestor & AST Parsing (`ingestor/parser.py`)
The parser leverages `tree-sitter` (specifically `tree-sitter-python`) to generate concrete syntax trees for Python modules. Instead of performing simple regex strings checks, it parses actual code scopes:
* **Function definitions**: Identifies function boundaries, names, and exact start/end line coordinates.
* **Class definitions**: Detects class definitions and measures scope length.
* **Imports statement**: Extracts imported packages and local module imports.
* **Calls**: Records internal function calls to map call trees and relationships.

### B. Vulnerability Scanner (`ingestor/vulnerability_scanner.py`)
This component parses `requirements.txt` to isolate package dependencies and their pinned versions.
1. The scanner triggers a `POST` request to `https://api.osv.dev/v1/query` containing the package names and versions.
2. The OSV API matches packages against the **PyPI ecosystem database** of open-source vulnerabilities.
3. Identified vulnerability objects (including CVE numbers, GitHub Security Advisories, summaries, and severity) are parsed.
4. The scanner merges them into Neo4j with a `[:HAS_VULNERABILITY]` relation linking back to the target dependency.

### C. Cognitive Security Auditor (Gemini Integration)
Under production mode, the `/assess-risk` endpoint acts as a cognitive security layer:
1. The backend performs a Cypher query on Neo4j for the target file or function, grabbing all related entities and connections (up to 50 nodes).
2. The resulting context (files, function callers, imports, active vulnerabilities) is packaged into a structured prompt.
3. The prompt is sent to Google's `gemini-1.5-flash` model.
4. Gemini returns an autonomous risk score (LOW/MEDIUM/HIGH/CRITICAL), a summary of architectural risk, possible attack vectors, and specific code-level mitigations.

---

## 3. Database Graph Schema
The Neo4j graph maps the target codebase's architectural relationships. 

### Nodes
* **`Project`**: Root node of the scanned repository.
* **`File`**: Represents a physical source code file in the repository.
* **`Class`**: Classes defined within a file.
* **`Function`**: Functions defined within a file.
* **`Dependency`**: Libraries imported by files or declared in requirements.
* **`Vulnerability`**: Known CVE/GHSA vulnerabilities mapped from OSV.

### Relationships
* **`Project`** - `[:HAS_FILE]` -> **`File`**
* **`Project`** - `[:DEPENDS_ON]` -> **`Dependency`**
* **`File`** - `[:CONTAINS]` -> **`Class`**
* **`File`** - `[:CONTAINS]` -> **`Function`**
* **`File`** - `[:IMPORTS]` -> **`Dependency`**
* **`File`** - `[:CALLS]` -> **`Function`**
* **`Dependency`** - `[:HAS_VULNERABILITY]` -> **`Vulnerability`**

---

## 4. API Documentation

The backend serves a FastAPI REST interface running on port `8000`.

| Method | Endpoint | Description |
| :--- | :--- | :--- |
| `GET` | `/` | Checks health status and reports active backend execution mode. |
| `GET` | `/health` | Diagnostic details showing status of Neo4j driver and Gemini configuration. |
| `GET` | `/graph/stats` | Fetches active node counts categorized by label from Neo4j. |
| `GET` | `/graph/vulnerabilities` | Returns parsed list of vulnerabilities and the files importing affected packages. |
| `DELETE` | `/graph/clear` | Purges graph entries associated with the specified `project_name` parameter. |
| `POST` | `/ingest` | Initiates asynchronous AST scanner & OSV vulnerability job in background. |
| `GET` | `/ingest/status` | Returns the current state and progress message. |
| `POST` | `/assess-risk` | Submits file/function entities to the AI cognitive auditor for security threat assessments. |

---

## 5. Configuration & Environment Variables

Create a `.env` file in the `backend/` and `ingestor/` folders to control server runtime behavior:

```env
NEO4J_URI=bolt://localhost:7687
NEO4J_USER=neo4j
NEO4J_PASSWORD=supersecretpassword
GEMINI_API_KEY=AIzaSy...
PORT=8000
ALLOWED_ORIGINS=http://localhost:5173,http://localhost:3000
```

---

## 6. How to Run the Application

The application supports a **Mock Simulation Mode** that allows complete offline validation without requiring a live Neo4j instance or Gemini credentials. If `NEO4J_URI=mock` is set, the application generates a complete simulation database in memory.

### Step 1: Initialize the Backend
1. Navigate to the backend directory:
   ```bash
   cd D:/projects/rishi/cns-mvp/backend
   ```
2. Install Python dependencies:
   ```bash
   pip install -r requirements.txt
   pip install -r ../ingestor/requirements.txt
   ```
3. Start the FastAPI development server:
   ```bash
   python main.py
   ```

### Step 2: Initialize the Frontend
1. Navigate to the frontend directory:
   ```bash
   cd D:/projects/rishi/cns-mvp/frontend
   ```
2. Install dependencies:
   ```bash
   npm install
   ```
3. Start the Vite dev server:
   ```bash
   npm run dev
   ```
"""
    },
    {
        "id": "bugbounty",
        "title": "BugBounty-AI (Recon Engine)",
        "docs_dir": "D:/projects/BugBoumty/docs",
        "filename": "bugbounty_technical_documentation",
        "content": """# BugBounty-AI — Technical Documentation
*Automated AI-Assisted Reconnaissance & Asset Discovery Engine*

This document provides a comprehensive technical overview of the **BugBounty-AI** project located at `D:/projects/BugBoumty`. It covers the system architecture, background worker queues, recon scanner pipeline, and integration endpoints.

---

## 1. Executive Summary
**BugBounty-AI** is a robust, distributed reconnaissance scanner built to automate defensive threat triage and attack surface mapping. It handles subdomain discovery, active port scanning, and web path crawling concurrently. Discovered assets are stored in a relational database ledger (PostgreSQL) and parsed by AI endpoints (OpenAI GPT or local Ollama instances) to highlight high-priority exploit paths.

---

## 2. System Architecture

The application is structured for distributed scalability:
* **`api/`**: FastAPI backend exposing REST endpoints for task triggers, targets, and result queries.
* **`workers/`**: Celery worker nodes tasked with running CLI recon utilities in parallel.
* **`db/`**: PostgreSQL database store mapping networks, ports, and live domains.
* **`ai/`**: Parsing adapters interfacing with OpenAI APIs or Ollama instances to evaluate logs.

```
┌─────────────────┐       Add Target       ┌─────────────────┐
│   Vite React    │ ─────────────────────► │  FastAPI Web    │
│    Dashboard    │                        │  Service (:8000)│
└─────────────────┘                        └────────┬────────┘
                                                    │ Celery Task Queue
                                                    ▼
┌─────────────────┐   Scan Target Assets   ┌─────────────────┐
│  PostgreSQL     │ ◄───────────────────── │  Celery Workers │
│  Ledger DB      │                        │  (with Redis)   │
└─────────────────┘                        └────────┬────────┘
                                                    │ Parse Logs via API
                                                    ▼
                                           ┌─────────────────┐
                                           │ AI Triager (LLM)│
                                           └─────────────────┘
```

### A. Distributed Recon Pipeline
The platform leverages standard offensive CLI utilities wrapped inside containerized Python execution environments:
1. **`subfinder`**: Discovers subdomains for the target origin via passive DNS queries.
2. **`httpx`**: Probes discovered domains to confirm online HTTP/HTTPS web servers, status codes, and titles.
3. **`katana`**: Crawls active web pages to identify input endpoints, forms, script sources, and API routes.

### B. Celery Task Queue & Redis
Scanning heavy domains is a blocking operation. FastAPI handles requests asynchronously by offloading scan execution to a **Celery Worker Pool** backed by a **Redis Broker**. Tasks are executed sequentially or in parallel depending on the scale of the worker pool, avoiding server CPU exhaustion during intensive scans.

---

## 3. Database Schema Layout
The PostgreSQL ledger stores the state of discovered network assets:
* **`Target`**: Main company domain or IP range defined by user.
* **`Subdomain`**: Subdomains resolved during subfinder scans.
* **`PortScan`**: Ports identified as open, with service name banners.
* **`CrawlRoute`**: Absolute URLs and parameters extracted by the crawling phase.
* **`AI_Flag`**: Custom alerts flagged by the LLM (potential vulnerability markers).

---

## 4. API Endpoints

The API is exposed via FastAPI on port `8000`.

| Method | Endpoint | Description |
| :--- | :--- | :--- |
| `GET` | `/` | Checks api operational status. |
| `POST` | `/targets/` | Adds a company domain to the scanning pool and queues a recon background task. |
| `GET` | `/targets/` | Returns list of all targets and overall scan states. |
| `GET` | `/targets/{id}/` | Detailed overview for a target. |
| `GET` | `/targets/{id}/subdomains` | Returns all subdomains discovered for the given target ID. |
| `GET` | `/targets/{id}/ports` | Returns list of open ports and associated services. |

---

## 5. Configuration & Setup

Create a `.env` file in the root directory:

```env
# Database Settings
POSTGRES_USER=postgres
POSTGRES_PASSWORD=dbpassword
POSTGRES_DB=bugbounty_ai
DATABASE_URL=postgresql://postgres:dbpassword@db:5432/bugbounty_ai

# Broker Settings
REDIS_URL=redis://redis:6379/0

# AI Configuration
OPENAI_API_KEY=sk-...
OLLAMA_HOST=http://localhost:11434
```

---

## 6. How to Run the Application

The entire platform is containerized using Docker and Docker Compose.

### Step 1: Deploy Services
1. Navigate to the project root:
   ```bash
   cd D:/projects/BugBoumty
   ```
2. Build and launch all containers:
   ```bash
   docker-compose up --build
   ```
   *This starts the API server, Celery worker pool, Redis, and PostgreSQL.*

### Step 2: Accessing the Systems
* **API Documentation**: [http://localhost:8000/docs](http://localhost:8000/docs) (Swagger UI)
* **Frontend Client**: [http://localhost:5173](http://localhost:5173) (if starting the Vite server locally inside the `frontend/` directory via `npm run dev`)
"""
    },
    {
        "id": "trading",
        "title": "Apex Quant (Trading Visualizer)",
        "docs_dir": "D:/projects/Trading/docs",
        "filename": "apex_quant_technical_documentation",
        "content": """# Apex Quant — Technical Documentation
*Interactive Quantitative Trading Visualizer & Strategy Engine*

This document provides a comprehensive technical overview of the **Apex Quant** platform located at `D:/projects/Trading`. It covers the trading strategy engine, real-time charting pipeline, API models, and setup guides.

---

## 1. Executive Summary
**Apex Quant** is an algorithmic trading visualizer and technical indicator platform. It calculations real-time pricing feeds using WebSockets, runs a technical indicator engine (EMA Ribbon + RSI Pullback), and plots live candlestick metrics, entry flags, stop-loss lines, and performance dashboards.

---

## 2. Trading Strategy Engine
The core trading algorithm implements the **EMA Ribbon + RSI Pullback** strategy.

### A. Trend highway (EMA Ribbon)
The engine maintains three Exponential Moving Averages (EMA):
1. **EMA 9 (Blue)**: Represents short-term momentum.
2. **EMA 21 (Amber)**: Represents intermediate trend direction.
3. **EMA 55 (Pink)**: Acts as the structural support or resistance level.

* **Bullish Alignment**: EMA 9 > EMA 21 > EMA 55 (Uptrend)
* **Bearish Alignment**: EMA 9 < EMA 21 < EMA 55 (Downtrend)

### B. Entry Rules & Execution
* **LONG Setup**: 
  * Trend is Bullish (9 > 21 > 55).
  * Price pulls back to touch or approach the 21/55 EMA ribbon lines.
  * RSI (14) drops below **40** (Momentum oversold).
* **SHORT Setup**: 
  * Trend is Bearish (9 < 21 < 55).
  * Price bounces back to touch or approach the 21/55 EMA ribbon lines.
  * RSI (14) rises above **60** (Momentum overbought).
* **Risk Management**: Enforces a strict **1:2.5 Risk-to-Reward Ratio**. The Stop Loss is placed slightly past the EMA 55 ribbon line.

```
Trend Highway Alignment (9, 21, 55 EMAs)
       │
       ├──► Bullish (9 > 21 > 55) ──► RSI < 40? ──► LONG Trade Entry (1:2.5 R:R)
       │
       └──► Bearish (9 < 21 < 55) ──► RSI > 60? ──► SHORT Trade Entry (1:2.5 R:R)
```

---

## 3. Technology Stack

* **Frontend**: Next.js (React 19), Zustand (State Management), TradingView Lightweight Charts (High-performance Canvas Rendering).
* **Backend**: Express.js, Node.js, WebSockets (streaming price feeds).
* **Database**: PostgreSQL (Prisma ORM for tracking transactions and logs).

---

## 4. API & WebSocket Architecture

### WebSocket Feeds (Backend to Frontend)
The Express backend streams match feeds on interval:
```json
{
  "symbol": "BTCUSDT",
  "price": 95243.20,
  "timestamp": 1719875420000,
  "indicators": {
    "ema9": 95120.50,
    "ema21": 95080.10,
    "ema55": 94910.00,
    "rsi": 38.5
  },
  "signal": "LONG"
}
```

### Express REST Endpoints
* `GET /api/history`: Retrieves database transaction objects of past trades.
* `POST /api/settings`: Updates active strategy thresholds (e.g. changing RSI bands).
* `GET /api/status`: Health metrics of pricing feed connections.

---

## 5. Configuration & Run Guide

Create a `.env` file inside the `backend/` folder:

```env
PORT=4000
DATABASE_URL=postgresql://user:password@localhost:5432/trading_db
MOCK_FEED=true
```

### Run Backend
1. Go to backend:
   ```bash
   cd D:/projects/Trading/backend
   ```
2. Install and launch:
   ```bash
   npm install
   npx prisma db push
   npm run dev
   ```

### Run Frontend
1. Go to frontend:
   ```bash
   cd D:/projects/Trading/frontend
   ```
2. Install and launch:
   ```bash
   npm install
   npm run dev
   ```
   *Open [http://localhost:3000](http://localhost:3000) to view the client.*
"""
    },
    {
        "id": "safetracker",
        "title": "SafeTracker (GPS App)",
        "docs_dir": "E:/PROJECTS/INTERN/locaton_tracking/docs",
        "filename": "safetracker_technical_documentation",
        "content": """# SafeTracker — Technical Documentation
*Consent-First Real-Time GPS Tracking Ecosystem*

This document provides a comprehensive technical overview of the **SafeTracker** location mapping system located at `E:/PROJECTS/INTERN/locaton_tracking`. It details the ethical tracking frameworks, React Native app, Node.js backend, and map interfaces.

---

## 1. Executive Summary
**SafeTracker** is an ethical, consent-first location-tracking application. It features a React Native client app for transmitting background coordinates, an Express.js server hosted on Render.com, and a real-time monitor web dashboard utilizing Leaflet.js to plot locations, accuracies, and paths.

---

## 2. Privacy & Ethical Guidelines
This platform is engineered strictly for **consensual** tracking (e.g. family location sharing, logistics). It enforces privacy via:
* **Explicit Opt-in**: The app requires the user to read and check a consent screen.
* **Persistent Notification Bar**: A permanent Android notification `📍 SafeTracker Active` is displayed whenever tracking is active.
* **Granular Stop Switch**: Users can deactivate tracking directly within the mobile client at any time, instantly halting background service tasks.
* **Device ID Visibility**: The hardware device ID is displayed inside the app header for total visibility.

---

## 3. Technology Stack & Directory Structure

```
locaton_tracking/
├── server/            ← Express.js + WebSocket server
├── dashboard/         ← Leaflet.js web monitoring client (HTML/JS)
└── SafeTracker/       ← React Native Mobile App (Expo SDK 54)
```

* **Mobile Client**: React Native, Expo SDK 54, `expo-location`, `expo-task-manager` (handles OS background location threads), and AsyncStorage (offline database buffering).
* **Backend Cloud**: Node.js, Express, `ws` (WebSockets).
* **Web Dashboard**: Vanilla HTML5/JS, Leaflet.js (map rendering), CartoDB dark-themed tile vectors.

---

## 4. Technical Features

### A. Offline Storage & Auto-Sync
If internet drops:
1. Coordinates are recorded by the background task.
2. Coordinates are stored as JSON arrays in `AsyncStorage` (buffers up to 2000 coordinates, ~5.5 hours of tracking).
3. The app displays an amber status banner: `X locations queued offline`.
4. Once internet returns, NetInfo triggers a `POST` request to `/api/location/batch` to upload coordinates in a single optimized payload.

### B. Hardware Device ID Resolution
To support multi-device dashboards, each device resolves a unique ID:
1. `androidId` (Android) or `identifierForVendor` (iOS) is pulled.
2. Stored in AsyncStorage and cached in memory.
3. Multiple phones appear as separate cards on the central map.

---

## 5. Endpoints & Run Guide

### API Endpoints
* `POST /api/location`: Records a single coordinate.
* `POST /api/location/batch`: Ingests queued offline coordinates.
* `GET /api/devices`: Lists all active unique device IDs.
* `GET /api/location/:deviceId/history`: Retrieves routing paths for a target device.

### How to Start Server
```bash
cd E:/PROJECTS/INTERN/locaton_tracking/server
npm install
node server.js
```
*Server runs on port 4000.*

### How to Start Expo App
```bash
cd E:/PROJECTS/INTERN/locaton_tracking/SafeTracker
npm install
npx expo start
```
"""
    },
    {
        "id": "iot_guardian",
        "title": "IoT Guardian Pro (Breach Sim)",
        "docs_dir": "E:/PROJECTS/INTERN/PROJECT_ONE/breach-simu/docs",
        "filename": "iot_guardian_technical_documentation",
        "content": """# IoT Guardian Pro — Technical Documentation
*Interactive Security Breach Simulator & Network Analytics Dashboard*

This document provides a comprehensive technical overview of the **IoT Guardian Pro** simulation system located at `E:/PROJECTS/INTERN/PROJECT_ONE/breach-simu`.

---

## 1. Executive Summary
**IoT Guardian Pro** is an educational cyber-range and security simulator. It runs automated, harmless attack simulations (DDoS spikes, Dictionary Brute-Force, and Phishing) to teach administrators how anomalies look on network charts. It supports live telemetry via WebSockets and exports PDF/Excel incident reports.

---

## 2. Platform Architecture

The project has four distinct components:
1. **FastAPI Backend (`backend/`)**: Handles database entries, logs match history, manages SQLite states, and executes WebSocket network simulations.
2. **React Dashboard (`frontend/`)**: Displays real-time packet speeds, attack vectors, system stats, and generates audit reports.
3. **Landing Page (`landing-page/`)**: Introducing corporate portal services.
4. **Phishing Target (`phishing-site/`)**: A mock login screen (`login.html`) to demonstrate phishing vector captures.

```
                  ┌──────────────────────┐
                  │    React Dashboard   │
                  └──────────┬───────────┘
                             │ HTTP / WebSockets
                             ▼
                  ┌──────────────────────┐
                  │    FastAPI Backend   │
                  └──────┬────────┬──────┘
                         │        │
            ┌────────────┘        └────────────┐
            ▼                                  ▼
   ┌──────────────────┐               ┌──────────────────┐
   │ SQLite Database  │               │ Simulations      │
   │ - Audit Logs     │               │ - DDoS Floods    │
   │ - User Records   │               │ - Brute Force    │
   └──────────────────┘               │ - Phishing Site  │
                                      └──────────────────┘
```

---

## 3. Core Simulation Modules

### A. DDoS Simulation
Simulates packet spikes. When activated:
* The backend increases simulated packets/second rates by 500%.
* Sends high-speed streams via `/ws` to the frontend.
* Charts display latency spikes and packet rate peaks.

### B. Brute-Force Simulation
Runs a mock login dictionary attack:
* Ingests passwords from a trimmed `rockyou.txt` dictionary file.
* Simulates hundreds of invalid login attempts against the system.
* Triggers alert notifications once threshold rates are breached.

### C. Phishing Simulation
* Hosts a mock authentication page.
* Records username/password credentials entered during simulation runs (safely logged inside SQLite).
* Teaches users about credentials harvesting vectors.

---

## 4. Reports Generation (PDF & Excel)

The dashboard allows exporting audit sheets:
* **PDF Audits**: Uses `reportlab` to design dynamic PDFs detailing incident timelines, attack categories, IP origins, and severity.
* **Excel Reports**: Uses `openpyxl` to write data rows of all simulated logs, allowing easy parsing by analyst teams.

---

## 5. Setup & Execution Instructions

### Setup Backend
1. Go to backend:
   ```bash
   cd E:/PROJECTS/INTERN/PROJECT_ONE/breach-simu/backend
   ```
2. Install libraries:
   ```bash
   pip install -r requirements.txt
   ```
3. Run FastAPI:
   ```bash
   uvicorn app.main:app --reload
   ```

### Setup Frontend
1. Go to frontend:
   ```bash
   cd E:/PROJECTS/INTERN/PROJECT_ONE/breach-simu/frontend
   ```
2. Install and launch:
   ```bash
   npm install
   npm run dev
   ```
"""
    },
    {
        "id": "nexalith",
        "title": "Nexalith Prime (Corporate Portal)",
        "docs_dir": "D:/projects/RoBlockSec/docs",
        "filename": "nexalith_technical_documentation",
        "content": """# Nexalith Prime — Technical Documentation
*Futuristic Corporate Portal & Security Agency Showcase*

This document provides a comprehensive technical overview of the **Nexalith Prime** corporate portal located at `D:/projects/RoBlockSec`.

---

## 1. Executive Summary
**Nexalith Prime** is a commercial portal designed for a cybersecurity agency (RoBlockSec). Built using React, Vite, and TailwindCSS, it displays interactive service descriptions (GRC, Red Teaming, Pen testing), careers portals, and downloadable product details. It features page transitions (Framer Motion), interactive particle backgrounds (tsparticles), and SEO-optimized routing (React Helmet Async).

---

## 2. Directory Layout & Key Modules

* **`src/pages/`**: Includes sections for:
  * `HomePage.tsx`: General landing with service showcases and call-to-actions.
  * `ServicesPage.tsx`: Breakdown of Red Teaming, SOC monitoring, and GRC consulting.
  * `ProductsPage.tsx`: Downloadable technical whitepapers and specifications.
  * `CareersPage.tsx`: Active jobs (SOC Analyst, Pentester) and application portals.
  * `TeamPage.tsx`, `AboutPage.tsx`, `BlogPage.tsx`: Administrative profiles and research.
* **`src/components/`**: Custom reusable UI components, headers, footers, and cards.
* **`src/constants/`**: Hardcoded texts and lists.

---

## 3. Frontend Technology & Libraries

* **Framework**: React 18, Vite.
* **Visual Styling**: TailwindCSS (responsive layouts, custom dark backgrounds).
* **Animations**: Framer Motion (fade-ins, scroll-triggered card translations).
* **Interactive Backdrop**: `tsparticles` (custom floating network nodes and nodes-linking lines).
* **SEO Metadata**: `react-helmet-async` (injects custom page meta headers dynamically for search indexers).
* **Routing**: React Router DOM (client-side routing).

---

## 4. How to Run the Application

To install dependencies and start the local development server:

### Step 1: Install Node Dependencies
1. Navigate to the project folder:
   ```bash
   cd D:/projects/RoBlockSec
   ```
2. Install npm packages:
   ```bash
   npm install
   ```

### Step 2: Start Vite Dev Server
1. Launch the server:
   ```bash
   npm run dev
   ```
2. Open the page:
   * **Local Address**: [http://localhost:5173](http://localhost:5173)

### Step 3: Build for Production
To generate a static distribution folder:
```bash
npm run build
```
*Creates a `dist/` directory ready for deployment on Netlify, Vercel, or AWS S3.*
"""
    }
]

# Formatting function for docx helper
def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

# Formatting functions for pdf xml clean
def clean_xml(text):
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'`(.*?)`', r'<font name="Courier" color="#C00000"><b>\1</b></font>', text)
    return text

# Compile single project function
def compile_project_docs(proj):
    title = proj["title"]
    docs_dir = proj["docs_dir"]
    filename = proj["filename"]
    md_content = proj["content"]

    print(f"\n=======================================================")
    print(f" Processing: {title}")
    print(f"=======================================================")

    # Ensure docs directory exists
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir, exist_ok=True)
        print(f"Created directory: {docs_dir}")

    # Write Markdown file
    md_file_path = os.path.join(docs_dir, f"{filename}.md")
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Saved MD file to: {md_file_path}")

    # Write DOCX
    docx_file_path = os.path.join(docs_dir, f"{filename}.docx")
    print(f"Compiling DOCX to: {docx_file_path}")
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_docx_heading(text, level):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            pPr = h._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1F4E78"/></w:pBdr>')
            pPr.append(pBdr)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        return h

    lines = md_content.splitlines()
    in_code = False
    code_lines = []
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()
        
        # Code Block
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_after = Pt(6)
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="7F7F7F"/></w:pBdr>')
                p._p.get_or_add_pPr().append(pBdr)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                p._p.get_or_add_pPr().append(shading)
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                code_lines = []
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Tables
        if stripped.startswith("|"):
            if ":" in stripped or "---" in stripped:
                continue
            in_table = True
            table_data.append([cell.strip() for cell in stripped.split("|")[1:-1]])
            continue
        else:
            if in_table and len(table_data) > 0:
                nr = len(table_data)
                nc = len(table_data[0]) if nr > 0 else 0
                if nc > 0:
                    t = doc.add_table(rows=nr, cols=nc)
                    t.style = 'Light Shading Accent 1'
                    t.autofit = True
                    for r_idx, row in enumerate(t.rows):
                        for c_idx, cell in enumerate(row.cells):
                            cell.text = table_data[r_idx][c_idx]
                            if r_idx == 0:
                                set_cell_background(cell, "1F4E78")
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.bold = True
                                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "F2F2F2")
                table_data = []
                in_table = False

        if not stripped:
            continue

        if stripped.startswith("# "):
            add_docx_heading(stripped[2:], 1)
        elif stripped.startswith("## "):
            add_docx_heading(stripped[3:], 2)
        elif stripped.startswith("### "):
            add_docx_heading(stripped[4:], 3)
        elif stripped.startswith("* ") or stripped.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            bold_parts = re.split(r'(\*\*.*?\*\*)', stripped[2:])
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        else:
            if stripped == "---":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                p.add_run("____________________________________________________").font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
                continue
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            bold_parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    code_parts = re.split(r'(`.*?`)', part)
                    for c_part in code_parts:
                        if c_part.startswith("`") and c_part.endswith("`"):
                            cr = p.add_run(c_part[1:-1])
                            cr.font.name = 'Consolas'
                            cr.font.size = Pt(9.5)
                            cr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        else:
                            p.add_run(c_part)

    # Footer
    f_p = sections[0].footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_p.add_run("Page ").font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    add_page_number(f_p.add_run())
    doc.save(docx_file_path)

    # Write PDF
    pdf_file_path = os.path.join(docs_dir, f"{filename}.pdf")
    print(f"Compiling PDF to: {pdf_file_path}")
    pdf = SimpleDocTemplate(
        pdf_file_path,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )

    styles = getSampleStyleSheet()
    pdf_h1 = ParagraphStyle('H1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=20, leading=24, textColor=colors.HexColor('#1F4E78'), spaceAfter=12, spaceBefore=14, keepWithNext=True)
    pdf_h2 = ParagraphStyle('H2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, leading=18, textColor=colors.HexColor('#2E74B5'), spaceAfter=8, spaceBefore=12, keepWithNext=True)
    pdf_h3 = ParagraphStyle('H3', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#595959'), spaceAfter=6, spaceBefore=10, keepWithNext=True)
    pdf_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#333333'), spaceAfter=8)
    pdf_list = ParagraphStyle('List', parent=styles['Normal'], fontName='Helvetica', fontSize=10, leading=14, textColor=colors.HexColor('#333333'), leftIndent=20, spaceAfter=4)
    pdf_code = ParagraphStyle('Code', parent=styles['Normal'], fontName='Courier', fontSize=8.5, leading=11, textColor=colors.HexColor('#404040'), leftIndent=15, spaceAfter=8)
    pdf_table_text = ParagraphStyle('TableText', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12, textColor=colors.HexColor('#333333'))
    pdf_table_header = ParagraphStyle('TableHeader', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9, leading=12, textColor=colors.white)

    story = []
    in_code = False
    code_lines = []
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()
        
        # Code Block
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                p_code = Paragraph(f"<pre>{clean_xml(chr(10).join(code_lines))}</pre>", pdf_code)
                t = Table([[p_code]], colWidths=[500])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F2F2F2')),
                    ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#7F7F7F')),
                    ('LEFTPADDING', (0,0), (-1,-1), 10),
                    ('RIGHTPADDING', (0,0), (-1,-1), 10),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ]))
                story.append(t)
                story.append(Spacer(1, 8))
                code_lines = []
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        # Tables
        if stripped.startswith("|"):
            if ":" in stripped or "---" in stripped:
                continue
            in_table = True
            table_data.append([cell.strip() for cell in stripped.split("|")[1:-1]])
            continue
        else:
            if in_table and len(table_data) > 0:
                nr = len(table_data)
                nc = len(table_data[0]) if nr > 0 else 0
                if nc > 0:
                    formatted_table_data = []
                    for r_idx, r_cells in enumerate(table_data):
                        formatted_row = []
                        for cell in r_cells:
                            c_clean = clean_xml(cell)
                            if r_idx == 0:
                                formatted_row.append(Paragraph(c_clean, pdf_table_header))
                            else:
                                formatted_row.append(Paragraph(c_clean, pdf_table_text))
                        formatted_table_data.append(formatted_row)
                    cw = 500 / nc
                    t = Table(formatted_table_data, colWidths=[cw]*nc)
                    t_styles = [
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E78')),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D0D0D0')),
                        ('TOPPADDING', (0,0), (-1,-1), 5),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                    ]
                    for r_idx in range(1, nr):
                        if r_idx % 2 == 0:
                            t_styles.append(('BACKGROUND', (0, r_idx), (-1, r_idx), colors.HexColor('#F2F2F2')))
                    t.setStyle(TableStyle(t_styles))
                    story.append(t)
                    story.append(Spacer(1, 8))
                table_data = []
                in_table = False

        if not stripped:
            continue

        if stripped.startswith("# "):
            story.append(Paragraph(clean_xml(stripped[2:]), pdf_h1))
        elif stripped.startswith("## "):
            story.append(Paragraph(clean_xml(stripped[3:]), pdf_h2))
        elif stripped.startswith("### "):
            story.append(Paragraph(clean_xml(stripped[4:]), pdf_h3))
        elif stripped.startswith("* ") or stripped.startswith("- "):
            story.append(Paragraph("• " + clean_xml(stripped[2:]), pdf_list))
        else:
            if stripped == "---":
                story.append(Spacer(1, 10))
                hr = Table([['']], colWidths=[500])
                hr.setStyle(TableStyle([
                    ('LINEBELOW', (0,0), (-1,-1), 1, colors.HexColor('#D0D0D0')),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0),
                ]))
                story.append(hr)
                story.append(Spacer(1, 10))
                continue
            story.append(Paragraph(clean_xml(line), pdf_body))

    def add_pdf_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(colors.HexColor('#7F7F7F'))
        canvas.drawRightString(612 - 54, 36, f"Page {doc.page}")
        canvas.restoreState()

    pdf.build(story, onFirstPage=add_pdf_footer, onLaterPages=add_pdf_footer)
    print(f"Saved PDF file to: {pdf_file_path}")

# Run compilation for all defined projects
if __name__ == "__main__":
    for proj in projects:
        try:
            compile_project_docs(proj)
        except Exception as e:
            print(f"[ERROR] Failed to compile docs for {proj['title']}: {e}")
    print("\nAll projects documents compiled in all 3 formats successfully!")
