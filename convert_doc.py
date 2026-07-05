import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# File paths
md_path = "cns_technical_documentation.md"
docx_path = "cns_technical_documentation.docx"
pdf_path = "cns_technical_documentation.pdf"

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

# ---------------------------------------------------------------------------
# 1. READ MD FILE
# ---------------------------------------------------------------------------
with open(md_path, "r", encoding="utf-8") as f:
    md_content = f.read()

# Split into lines
lines = md_content.splitlines()

# ---------------------------------------------------------------------------
# 2. GENERATE DOCX FILE
# ---------------------------------------------------------------------------
print("Generating DOCX file...")
doc = Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Default style fonts
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper to add headings
def add_custom_heading(text, level, space_before=12, space_after=6):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    run = h.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        # Add bottom border XML to H1 for styling
        pPr = h._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1F4E78"/></w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return h

# Parse lines and write DOCX
in_code = False
code_lines = []
in_table = False
table_data = []

# Simple Markdown Parsing Loop for DOCX
for line in lines:
    stripped = line.strip()
    
    # 1. Code Blocks
    if stripped.startswith("```"):
        if in_code:
            in_code = False
            # Write code block paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="7F7F7F"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
            p._p.get_or_add_pPr().append(shading)
            
            code_text = "\n".join(code_lines)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            code_lines = []
        else:
            in_code = True
        continue
        
    if in_code:
        code_lines.append(line)
        continue

    # 2. Tables
    if stripped.startswith("|"):
        # Check if separator line
        if ":" in stripped or "---" in stripped:
            continue
        in_table = True
        # Parse table row
        row_cells = [cell.strip() for cell in stripped.split("|")[1:-1]]
        table_data.append(row_cells)
        continue
    else:
        if in_table and len(table_data) > 0:
            # Build Table in docx
            num_rows = len(table_data)
            num_cols = len(table_data[0]) if num_rows > 0 else 0
            if num_cols > 0:
                t = doc.add_table(rows=num_rows, cols=num_cols)
                t.style = 'Light Shading Accent 1'
                t.autofit = True
                
                # Apply background shading and style
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell.text = table_data[r_idx][c_idx]
                        # Style headers
                        if r_idx == 0:
                            set_cell_background(cell, "1F4E78")
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            if r_idx % 2 == 0:
                                set_cell_background(cell, "F2F2F2")
            table_data = []
            in_table = False

    # Skip empty lines
    if not stripped:
        continue

    # 3. Headings
    if stripped.startswith("# "):
        add_custom_heading(stripped[2:], 1, space_before=18, space_after=10)
    elif stripped.startswith("## "):
        add_custom_heading(stripped[3:], 2, space_before=14, space_after=8)
    elif stripped.startswith("### "):
        add_custom_heading(stripped[4:], 3, space_before=12, space_after=6)
    
    # 4. Bullet lists
    elif stripped.startswith("* ") or stripped.startswith("- "):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        # Parse bold formatting in lists
        text_content = stripped[2:]
        bold_parts = re.split(r'(\*\*.*?\*\*)', text_content)
        for part in bold_parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
            else:
                p.add_run(part)
                
    # 5. Normal Paragraph
    else:
        # Check if separator line
        if stripped == "---":
            # Add section break or horizontal line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.add_run("____________________________________________________").font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        # Parse bold formatting like **text**
        bold_parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in bold_parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
            else:
                # Also handle inline code block style `code`
                code_parts = re.split(r'(`.*?`)', part)
                for c_part in code_parts:
                    if c_part.startswith("`") and c_part.endswith("`"):
                        cr = p.add_run(c_part[1:-1])
                        cr.font.name = 'Consolas'
                        cr.font.size = Pt(9.5)
                        cr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Dark red for inline code
                    else:
                        p.add_run(c_part)

# Add footer page numbers to DOCX
footer = sections[0].footer
f_p = footer.paragraphs[0]
f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
f_run = f_p.add_run("Page ")
f_run.font.size = Pt(9)
f_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
add_page_number(f_p.add_run())

doc.save(docx_path)
print(f"DOCX created successfully at: {docx_path}")

# ---------------------------------------------------------------------------
# 3. GENERATE PDF FILE USING REPORTLAB
# ---------------------------------------------------------------------------
print("Generating PDF file...")
pdf = SimpleDocTemplate(
    pdf_path,
    pagesize=letter,
    rightMargin=54, leftMargin=54,
    topMargin=54, bottomMargin=54
)

styles = getSampleStyleSheet()

# Custom styles
pdf_h1 = ParagraphStyle(
    'PdfH1',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=22,
    leading=26,
    textColor=colors.HexColor('#1F4E78'),
    spaceAfter=15,
    spaceBefore=18,
    keepWithNext=True
)

pdf_h2 = ParagraphStyle(
    'PdfH2',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=15,
    leading=19,
    textColor=colors.HexColor('#2E74B5'),
    spaceAfter=10,
    spaceBefore=14,
    keepWithNext=True
)

pdf_h3 = ParagraphStyle(
    'PdfH3',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=11,
    leading=14,
    textColor=colors.HexColor('#595959'),
    spaceAfter=8,
    spaceBefore=12,
    keepWithNext=True
)

pdf_body = ParagraphStyle(
    'PdfBody',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=10,
    leading=14,
    textColor=colors.HexColor('#333333'),
    spaceAfter=8
)

pdf_list = ParagraphStyle(
    'PdfList',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=10,
    leading=14,
    textColor=colors.HexColor('#333333'),
    leftIndent=20,
    spaceAfter=4
)

pdf_code = ParagraphStyle(
    'PdfCode',
    parent=styles['Normal'],
    fontName='Courier',
    fontSize=8.5,
    leading=11,
    textColor=colors.HexColor('#404040'),
    leftIndent=15,
    spaceAfter=8
)

pdf_table_text = ParagraphStyle(
    'PdfTableText',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=9,
    leading=12,
    textColor=colors.HexColor('#333333')
)

pdf_table_header = ParagraphStyle(
    'PdfTableHeader',
    parent=styles['Normal'],
    fontName='Helvetica-Bold',
    fontSize=9,
    leading=12,
    textColor=colors.white
)

story = []

def clean_xml(text):
    # Escape XML characters
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Convert MD bold **text** to <b>text</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Convert MD inline code `code` to <font name="Courier">code</font>
    text = re.sub(r'`(.*?)`', r'<font name="Courier" color="#C00000"><b>\1</b></font>', text)
    return text

# Reset state
in_code = False
code_lines = []
in_table = False
table_data = []

for line in lines:
    stripped = line.strip()
    
    # 1. Code Blocks
    if stripped.startswith("```"):
        if in_code:
            in_code = False
            code_text = "\n".join(code_lines)
            # Escape code block content
            code_text_clean = clean_xml(code_text)
            # We use preformatted Paragraph
            p = Paragraph(f"<pre>{code_text_clean}</pre>", pdf_code)
            # Put code block in a small single-cell table for background color and border
            t = Table([[p]], colWidths=[500])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F2F2F2')),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#7F7F7F')),
                ('LEFTPADDING', (0,0), (-1,-1), 10),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))
            code_lines = []
        else:
            in_code = True
        continue
        
    if in_code:
        code_lines.append(line)
        continue

    # 2. Tables
    if stripped.startswith("|"):
        if ":" in stripped or "---" in stripped:
            continue
        in_table = True
        row_cells = [cell.strip() for cell in stripped.split("|")[1:-1]]
        table_data.append(row_cells)
        continue
    else:
        if in_table and len(table_data) > 0:
            num_rows = len(table_data)
            num_cols = len(table_data[0]) if num_rows > 0 else 0
            if num_cols > 0:
                # Convert raw text to Paragraph flowables for cell auto-wrapping
                formatted_table_data = []
                for r_idx, row in enumerate(table_data):
                    formatted_row = []
                    for cell in row:
                        cell_clean = clean_xml(cell)
                        if r_idx == 0:
                            formatted_row.append(Paragraph(cell_clean, pdf_table_header))
                        else:
                            formatted_row.append(Paragraph(cell_clean, pdf_table_text))
                    formatted_table_data.append(formatted_row)
                
                # Setup column widths (distribute 500pt across columns)
                col_width = 500 / num_cols
                t = Table(formatted_table_data, colWidths=[col_width]*num_cols)
                
                # Apply styling
                t_styles = [
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E78')),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D0D0D0')),
                    ('TOPPADDING', (0,0), (-1,-1), 5),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ]
                
                # Alternating row colors
                for r_idx in range(1, num_rows):
                    if r_idx % 2 == 0:
                        t_styles.append(('BACKGROUND', (0, r_idx), (-1, r_idx), colors.HexColor('#F2F2F2')))
                        
                t.setStyle(TableStyle(t_styles))
                story.append(t)
                story.append(Spacer(1, 8))
            table_data = []
            in_table = False

    if not stripped:
        continue

    # 3. Headings
    if stripped.startswith("# "):
        story.append(Paragraph(clean_xml(stripped[2:]), pdf_h1))
    elif stripped.startswith("## "):
        story.append(Paragraph(clean_xml(stripped[3:]), pdf_h2))
    elif stripped.startswith("### "):
        story.append(Paragraph(clean_xml(stripped[4:]), pdf_h3))
    
    # 4. Bullet lists
    elif stripped.startswith("* ") or stripped.startswith("- "):
        bullet_item = "• " + clean_xml(stripped[2:])
        story.append(Paragraph(bullet_item, pdf_list))
        
    # 5. Normal Paragraph
    else:
        if stripped == "---":
            # Add spacing or horizontal rule lines
            story.append(Spacer(1, 10))
            hr = Table([['']], colWidths=[500])
            hr.setStyle(TableStyle([
                ('LINEBELOW', (0,0), (-1,-1), 1, colors.HexColor('#D0D0D0')),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(hr)
            story.append(Spacer(1, 10))
            continue
            
        p_text = clean_xml(line)
        story.append(Paragraph(p_text, pdf_body))

# Footer page callback
def add_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 9)
    canvas.setFillColor(colors.HexColor('#7F7F7F'))
    canvas.drawRightString(612 - 54, 36, f"Page {doc.page}")
    canvas.restoreState()

# Build document
pdf.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
print(f"PDF created successfully at: {pdf_path}")

print("Success! Created DOCX and PDF documents.")
